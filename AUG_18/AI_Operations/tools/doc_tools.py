import os
from langchain_core.tools import tool
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document

WORKSPACE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'workspace'))

def _get_safe_path(filename: str) -> str:
    target_path = os.path.abspath(os.path.join(WORKSPACE_DIR, filename))
    if not target_path.startswith(WORKSPACE_DIR):
        raise ValueError(f"Access to '{filename}' denied.")
    return target_path

@tool
def generate_pdf_report(filename: str, title: str, content: str) -> str:
    """Builds a styled multi-section PDF report."""
    try:
        path = _get_safe_path(filename)
        if not path.endswith('.pdf'):
            path += '.pdf'
            
        doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        
        for p in content.split('\n'):
            if p.strip():
                story.append(Paragraph(p, styles['Normal']))
                story.append(Spacer(1, 6))
                
        doc.build(story)
        return f"Successfully generated PDF: {os.path.basename(path)}"
    except Exception as e:
        return f"Error generating PDF: {e}"

@tool
def generate_excel_report(filename: str, sheet_title: str, headers: list[str], rows: list[list[str]]) -> str:
    """Builds a formatted Excel workbook with styled table columns. Pass headers and rows."""
    try:
        path = _get_safe_path(filename)
        if not path.endswith('.xlsx'):
            path += '.xlsx'
            
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_title
        
        # Write headers
        for col, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True)
            
        # Write rows
        for r_idx, row_data in enumerate(rows, start=2):
            for c_idx, val in enumerate(row_data, start=1):
                ws.cell(row=r_idx, column=c_idx, value=val)
                
        wb.save(path)
        return f"Successfully generated Excel: {os.path.basename(path)}"
    except Exception as e:
        return f"Error generating Excel: {e}"

@tool
def generate_word_doc(filename: str, title: str, content: str) -> str:
    """Formats executive briefings and operational memos into a Word doc."""
    try:
        path = _get_safe_path(filename)
        if not path.endswith('.docx'):
            path += '.docx'
            
        doc = Document()
        doc.add_heading(title, 0)
        
        for p in content.split('\n'):
            if p.strip():
                doc.add_paragraph(p)
                
        doc.save(path)
        return f"Successfully generated Word doc: {os.path.basename(path)}"
    except Exception as e:
        return f"Error generating Word doc: {e}"

doc_tools = [generate_pdf_report, generate_excel_report, generate_word_doc]
